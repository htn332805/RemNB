# Report generation using python-docx
import os, sqlite3
from docx import Document

def run(params):
    # Extract parameters from orchestrator
    template_path = params["template_path"]
    test_id = params["test_id"]
    db_path = params["db_path"]
    excel_path = params["excel_path"]
    ai_summary_path = params["ai_summary_path"]
    report_out = params["report_out"]

    # Load the template document
    try:
        doc = Document(template_path)
    except Exception:
        doc = Document()  # blank
        
    # Connect to SQLite to fetch test metadata and summary
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute("""SELECT tester, created_at, description, summary, issues, notes
                   FROM tests WHERE test_id=?""", (test_id,))
    row = cur.fetchone()
    conn.close()

    # Unpack DB row safely
    if row:
        tester, created_at, description, summary, issues, notes = row
    else:
        tester = created_at = description = summary = issues = notes = None

    # Add a report title
    doc.add_heading("Thermal Test Report", level=1)
    # Insert key metadata fields
    doc.add_paragraph(f"Test ID: {test_id}")
    doc.add_paragraph(f"Tester: {tester}")
    doc.add_paragraph(f"Created At: {created_at}")
    doc.add_paragraph(f"Description: {description}")

    # AI summary section
    doc.add_heading("AI Summary", level=2)
    if os.path.exists(ai_summary_path):
        # Read the markdown as plain text (simple approach)
        md_text = open(ai_summary_path, "r").read()
        for line in md_text.splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("No AI summary available.")

    # Artifacts section with Excel path
    doc.add_heading("Artifacts", level=2)
    doc.add_paragraph(f"Excel workbook: {excel_path}")

    # Ensure output directory exists
    os.makedirs(os.path.dirname(report_out), exist_ok=True)
    # Save the report document
    doc.save(report_out)

    # Return structured info for orchestrator logging
    return {"ok": True, "report": report_out}
